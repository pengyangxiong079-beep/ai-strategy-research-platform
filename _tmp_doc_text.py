from pathlib import Path

from docx import Document


folder = Path(r"C:\Users\leo\Desktop\1")
out = Path(r"C:\Users\leo\Desktop\ai-strategy-agent\_tmp_doc_text.txt")
lines = []
for path in sorted((p for p in folder.glob("*.docx") if not p.name.startswith("~$")), key=lambda p: p.stat().st_size):
    doc = Document(path)
    lines += ["=" * 80, path.name, f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}"]
    for i, p in enumerate(doc.paragraphs):
        lines.append(f"P{i:03d} [{p.style.name}] {p.text}")
    for ti, table in enumerate(doc.tables):
        lines.append(f"TABLE {ti} ({len(table.rows)}x{len(table.columns)})")
        for ri, row in enumerate(table.rows):
            lines.append(f"  R{ri:02d}: " + " || ".join(c.text.replace("\n", " / ") for c in row.cells))
out.write_text("\n".join(lines), encoding="utf-8")
print(out)

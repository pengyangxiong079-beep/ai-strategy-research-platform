from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def run_text(run):
    rpr = run._element.rPr
    fonts = {}
    if rpr is not None and rpr.rFonts is not None:
        fonts = {
            key: rpr.rFonts.get(qn(f"w:{key}"))
            for key in ("ascii", "hAnsi", "eastAsia", "cs")
        }
    return {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "underline": bool(run.underline) if run.underline is not None else None,
        "size_pt": run.font.size.pt if run.font.size else None,
        "font_name": run.font.name,
        "fonts": fonts,
    }


def paragraph_data(paragraph, index, location="body"):
    fmt = paragraph.paragraph_format
    return {
        "location": location,
        "index": index,
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text,
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
        "left_indent_pt": fmt.left_indent.pt if fmt.left_indent else None,
        "right_indent_pt": fmt.right_indent.pt if fmt.right_indent else None,
        "first_line_indent_pt": fmt.first_line_indent.pt if fmt.first_line_indent else None,
        "space_before_pt": fmt.space_before.pt if fmt.space_before else None,
        "space_after_pt": fmt.space_after.pt if fmt.space_after else None,
        "line_spacing": str(fmt.line_spacing) if fmt.line_spacing is not None else None,
        "page_break_before": fmt.page_break_before,
        "keep_with_next": fmt.keep_with_next,
        "runs": [run_text(run) for run in paragraph.runs],
    }


def main():
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    doc = Document(source)
    paragraphs = [paragraph_data(p, i) for i, p in enumerate(doc.paragraphs)]
    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for ri, row in enumerate(table.rows):
            cells = []
            for ci, cell in enumerate(row.cells):
                cells.append({
                    "column": ci,
                    "text": cell.text,
                    "paragraphs": [paragraph_data(p, pi, f"table:{ti}:{ri}:{ci}") for pi, p in enumerate(cell.paragraphs)],
                })
            rows.append(cells)
        tables.append({"index": ti, "style": table.style.name if table.style else None, "rows": rows})

    sections = []
    for si, section in enumerate(doc.sections):
        sections.append({
            "index": si,
            "width_in": section.page_width.inches,
            "height_in": section.page_height.inches,
            "top_margin_in": section.top_margin.inches,
            "bottom_margin_in": section.bottom_margin.inches,
            "left_margin_in": section.left_margin.inches,
            "right_margin_in": section.right_margin.inches,
            "header_distance_in": section.header_distance.inches,
            "footer_distance_in": section.footer_distance.inches,
            "header": [paragraph_data(p, i, f"header:{si}") for i, p in enumerate(section.header.paragraphs)],
            "footer": [paragraph_data(p, i, f"footer:{si}") for i, p in enumerate(section.footer.paragraphs)],
        })

    package = {}
    with zipfile.ZipFile(source) as zf:
        names = zf.namelist()
        package["parts"] = names
        for name in (
            "word/footnotes.xml",
            "word/endnotes.xml",
            "word/comments.xml",
            "word/numbering.xml",
            "word/styles.xml",
            "docProps/core.xml",
        ):
            if name in names:
                package[name] = zf.read(name).decode("utf-8", errors="replace")

    payload = {
        "source": str(source),
        "core_properties": {
            "title": doc.core_properties.title,
            "subject": doc.core_properties.subject,
            "author": doc.core_properties.author,
            "keywords": doc.core_properties.keywords,
            "comments": doc.core_properties.comments,
            "last_modified_by": doc.core_properties.last_modified_by,
            "created": str(doc.core_properties.created),
            "modified": str(doc.core_properties.modified),
        },
        "paragraphs": paragraphs,
        "tables": tables,
        "sections": sections,
        "inline_shapes": len(doc.inline_shapes),
        "package": package,
    }
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()

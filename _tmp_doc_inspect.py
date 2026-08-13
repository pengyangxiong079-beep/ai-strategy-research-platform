from __future__ import annotations

import json
from pathlib import Path

from docx import Document


def length_value(value):
    if value is None:
        return None
    return {"emu": int(value), "pt": value.pt, "cm": value.cm}


def paragraph_record(index, paragraph):
    fmt = paragraph.paragraph_format
    return {
        "index": index,
        "style": paragraph.style.name,
        "alignment": str(paragraph.alignment),
        "text": paragraph.text,
        "format": {
            "left_indent": length_value(fmt.left_indent),
            "right_indent": length_value(fmt.right_indent),
            "first_line_indent": length_value(fmt.first_line_indent),
            "space_before": length_value(fmt.space_before),
            "space_after": length_value(fmt.space_after),
            "line_spacing": str(fmt.line_spacing),
            "keep_with_next": fmt.keep_with_next,
            "page_break_before": fmt.page_break_before,
        },
        "runs": [
            {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": str(run.underline),
                "font": run.font.name,
                "size": length_value(run.font.size),
                "color": str(run.font.color.rgb) if run.font.color.rgb else None,
            }
            for run in paragraph.runs
        ],
    }


def inspect(path: Path):
    doc = Document(path)
    return {
        "path": str(path.resolve()),
        "size": path.stat().st_size,
        "sections": [
            {
                "page_width": length_value(s.page_width),
                "page_height": length_value(s.page_height),
                "top_margin": length_value(s.top_margin),
                "bottom_margin": length_value(s.bottom_margin),
                "left_margin": length_value(s.left_margin),
                "right_margin": length_value(s.right_margin),
                "header_distance": length_value(s.header_distance),
                "footer_distance": length_value(s.footer_distance),
                "start_type": str(s.start_type),
            }
            for s in doc.sections
        ],
        "paragraphs": [paragraph_record(i, p) for i, p in enumerate(doc.paragraphs)],
        "tables": [
            {
                "index": ti,
                "style": table.style.name if table.style else None,
                "rows": [
                    [
                        {
                            "text": cell.text,
                            "paragraphs": [paragraph_record(pi, p) for pi, p in enumerate(cell.paragraphs)],
                        }
                        for cell in row.cells
                    ]
                    for row in table.rows
                ],
            }
            for ti, table in enumerate(doc.tables)
        ],
        "inline_shapes": len(doc.inline_shapes),
        "styles": [
            {
                "name": style.name,
                "type": str(style.type),
                "font": style.font.name,
                "size": length_value(style.font.size),
                "bold": style.font.bold,
                "italic": style.font.italic,
            }
            for style in doc.styles
            if style.name in {p.style.name for p in doc.paragraphs}
        ],
    }


def main():
    folder = Path(r"C:\Users\leo\Desktop\1")
    output = Path(r"C:\Users\leo\Desktop\ai-strategy-agent\_tmp_doc_inspection.json")
    records = []
    for path in sorted(folder.glob("*.docx"), key=lambda p: p.stat().st_size):
        if path.name.startswith("~$"):
            continue
        try:
            records.append(inspect(path))
        except Exception as exc:
            records.append({"path": str(path), "error": repr(exc)})
    output.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    print(output)


if __name__ == "__main__":
    main()

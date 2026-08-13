from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE_DIR = Path(r"C:\Users\leo\Desktop\1")
WORK_DIR = Path(r"C:\Users\leo\Desktop\ai-strategy-agent")
TEMP_OUTPUT = WORK_DIR / "无生展示课_v2_一般过去时专业扩充版.docx"


PARAGRAPH_TEXT = {
    0: "无生展示课九年级一般过去时复习",
    1: "",
    2: "开场（0:00—0:40）：",
    3: "Class begins! Good morning, boys and girls! Sit down, please. Today we are going to open a special summer memory box and take a trip back in time.",
    4: "",
    5: "情境导入（0:40—1:20）：",
    6: "[微笑，展示模糊照片] Look! This photo is not clear, and some words are missing. Can you help me repair this summer memory?",
    7: "[停顿2秒，预设回应] Of course! / Yes, we can!",
    8: "First, is the event happening now, or did it happen in the past?",
    9: "[预设回应] It happened in the past. — Exactly. What clue helped you? [预设回应] Last summer.",
    10: "Great noticing! A time clue tells us when an action happened and helps us choose the correct tense.",
    11: "So here is today's big question: How can we make a past memory clear, correct and interesting?",
    12: "[指向时间轴] Put these expressions on the timeline: yesterday, last summer, two days ago and in 2025. Past or now? [停顿2秒] Yes, they all point to the past.",
    13: "Now compare the three sentences. In an affirmative sentence, use V-ed or an irregular past form. After didn't or Did, the verb returns to its base form. Say it together: did — base form. For be, choose was or were according to the subject.",
    14: "",
    15: "纠错探究（4:10—7:10）：",
    16: "Four memories are broken. Be language doctors: find each problem, repair it and give one reason. You have twenty seconds. Begin!",
    17: "[身体略后退，静默扫视；20秒后回到屏幕侧] Time's up. Let's check them one by one.",
    18: "Number one? Yes — went. You noticed the time clue. Number two? Excellent: didn't play. Number three, this side please... Right: Did you see...? Number four? The weather was sunny.",
    19: "",
    20: "归纳提升：",
    21: "Now compare sentences two and three. What do they have in common?",
    22: "[预设回应] Both use did, so both use the base form. — Exactly. Whether did is in a negative sentence or a question, the following verb stays in its base form.",
    23: "",
    24: "If you can find the time clue, repair the verb and explain why, you earn three memory stars: clue, form and reason.",
    25: "The sentences are correct now, but the memory is still incomplete. Let's interview the memory owner and add meaningful details.",
    26: "",
    27: "采访追问（7:10—11:20）：",
    28: "A clear memory answers four questions: Where did you go? Who did you go with? What did you do there? How was it? Notice that Did is followed by the base form, while answers use the past form.",
    29: "",
    30: "示范问答：",
    31: "Listen to my first question: Where did you go last summer?",
    32: "[预设回应] I went to Qingdao. — A complete answer, well done. Who did you go with?",
    33: "[预设回应] I went with my family. We walked along the beach and took photos. It was exciting. — I like the concrete details, especially walked along the beach.",
    34: "",
    35: "Now Student A is the memory detective and Student B is the memory owner. Ask at least two follow-up questions, then change roles. You have forty seconds. Go!",
    36: "[静默走两步，做倾听状；约25秒提示] Change roles, please. [40秒时击掌收回] Time's up.",
    37: "I heard a clear question from this pair: What did you do there? I also heard the detail we watched the sunrise. A useful follow-up question makes a memory clearer; a specific detail makes it more interesting.",
    38: "Remember: do not stop at one answer. Listen carefully and ask for one more detail.",
    39: "",
    40: "语篇输出（11:20—13:50）：",
    41: "Now turn your answers into a Memory Postcard. Write or say at least four sentences. Include one past-time clue, two accurate past verbs, one sequence word and one true feeling.",
    42: "Use the PAST checklist: P — Past time; A — Accurate verbs; S — Sequence; T — True feeling.",
    43: "For example: Last summer, I went to Qingdao with my family. First, we walked along the beach. Then, we took many photos. The trip was short but exciting, and I learned to enjoy time with my family.",
    44: "[指向作品区] Take thirty seconds to prepare your postcard. You may use the four questions and the PAST checklist as helpers.",
    45: "Exchange your postcards. Don't just say Good. Give one strength and one suggestion: Your time clue is clear. / Your verbs are accurate. / Add a sequence word. / Tell us more about your feeling.",
    46: "[展示反馈] Your time clue and past verbs are clear. If you add then before the third sentence, the story will be easier to follow. Specific feedback helps us improve.",
    47: "总结评价（13:50—15:00）： Before we close the memory box, complete the exit ticket. [依次停顿] Yesterday, Lucy ___ to the library. She didn't ___ TV. ___ she borrow a book? [预设回应] went; watch; Did. Excellent — you used all three core structures correctly.",
    48: "Today we learned to make a past memory clear: find a past-time clue, choose accurate verb forms, put events in sequence and add a true feeling. The past is not only a tense in a grammar book; it is the story of how we grew. Record it clearly and treasure it deeply. Homework: improve your Memory Postcard, or record a 60-second My Unforgettable Day. Thank you, everyone!",
}


TABLE_TEXT = [
    [
        "希沃1  A Trip Back in Time",
        "[模糊的暑假照片]   Memory loading...",
        "Is it happening now, or did it happen in the past?",
        "Time clue: last summer",
        "Big Question: How can we make a past memory clear, correct and interesting?",
    ],
    [
        "希沃2  Notice: Past-time map & sentence forms",
        "PAST: yesterday | last summer | two days ago | in 2025",
        "+  Last summer, I went to Qingdao.        S + V-ed / V2",
        "−  I didn't stay at home.                 S + didn't + V",
        "?  Did I travel alone?                    Did + S + V ...?",
        "be: I / He / She / It was ...             You / We / They were ...",
        "Key rule: did → base form",
    ],
    [
        "希沃3  Repair the broken memories",
        "1. I go to Qingdao last summer.       → I went to Qingdao last summer.",
        "2. We didn't played games.            → We didn't play games.",
        "3. Did you saw the sea?                → Did you see the sea?",
        "4. The weather were sunny.             → The weather was sunny.",
        "Challenge: Correct it + explain why.",
    ],
    [
        "希沃4  From correct sentences to a complete memory",
        "Where did you go?",
        "Who did you go with?",
        "What did you do there?",
        "How was it?",
        "Listen → answer in a full sentence → ask one follow-up question",
    ],
    [
        "希沃5  Memory Interview Card",
        "A: Where did you go last summer?",
        "B: I went to __________.",
        "A: Who did you go with? / What did you do there?",
        "B: I went with __________. We __________.",
        "A: How was it? Why?",
        "B: It was __________ because __________.",
    ],
    [
        "希沃6  PAST Self-check",
        "P · Past time      I used a clear past-time clue.",
        "A · Accurate verbs I used V2 or did + base form correctly.",
        "S · Sequence       I used first / then / after that.",
        "T · True feeling   I added a real feeling or reflection.",
    ],
    [
        "希沃7  Memory Postcard",
        "Last summer, I went to Qingdao with my family.",
        "First, we walked along the beach.",
        "Then, we took many photos.",
        "The trip was short but exciting, and I learned to enjoy time with my family.",
        "Peer feedback: one strength + one suggestion",
    ],
    [
        "希沃8  Exit Ticket & Closing",
        "Yesterday, Lucy ___ to the library.",
        "She didn't ___ TV.",
        "___ she borrow a book?",
        "PAST: Past time · Accurate verbs · Sequence · True feeling",
        "The past tells the story of how we grew.",
    ],
]


def first_run_properties(paragraph):
    for run in paragraph.runs:
        if run._r.rPr is not None:
            return deepcopy(run._r.rPr)
    return None


def replace_paragraph_text(paragraph, text, fallback_rpr=None):
    rpr = first_run_properties(paragraph) or deepcopy(fallback_rpr)
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    if text == "":
        return
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(rpr)
    node = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    paragraph._p.append(run)


def replace_cell_lines(cell, lines):
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.add_paragraph()
        paragraphs = list(cell.paragraphs)
    template = deepcopy(paragraphs[-1]._p)
    while len(paragraphs) < len(lines):
        new_p = deepcopy(template)
        cell._tc.append(new_p)
        paragraphs = list(cell.paragraphs)
    for paragraph, text in zip(paragraphs, lines):
        replace_paragraph_text(paragraph, text)
    for paragraph in paragraphs[len(lines):]:
        cell._tc.remove(paragraph._p)


def main():
    candidates = [p for p in SOURCE_DIR.glob("*.docx") if not p.name.startswith("~$")]
    source = min(candidates, key=lambda p: p.stat().st_size)
    shutil.copy2(source, TEMP_OUTPUT)
    doc = Document(TEMP_OUTPUT)
    if len(doc.paragraphs) != 49 or len(doc.tables) != 8:
        raise RuntimeError(f"Unexpected source structure: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    body_rpr = first_run_properties(doc.paragraphs[3])
    for index, text in PARAGRAPH_TEXT.items():
        replace_paragraph_text(doc.paragraphs[index], text, body_rpr)
    # Keep the two key instructional statements intact across page boundaries.
    doc.paragraphs[27].paragraph_format.keep_with_next = True
    doc.paragraphs[28].paragraph_format.keep_together = True
    doc.paragraphs[42].paragraph_format.keep_together = True
    for table, lines in zip(doc.tables, TABLE_TEXT):
        replace_cell_lines(table.cell(0, 0), lines)
    props = doc.core_properties
    props.title = "无生展示课九年级一般过去时复习（专业扩充版）"
    props.subject = "15分钟无生课堂教学逐字稿"
    props.keywords = "一般过去时, 无生课堂, 初中英语, 教学展示"
    doc.save(TEMP_OUTPUT)
    print(TEMP_OUTPUT)


if __name__ == "__main__":
    main()

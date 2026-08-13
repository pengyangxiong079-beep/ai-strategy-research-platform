from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


path = min((p for p in Path(r"C:\Users\leo\Desktop\1").glob("*.docx") if not p.name.startswith("~$")), key=lambda p: p.stat().st_size)
doc = Document(path)
p_lookup = {p._p: i for i, p in enumerate(doc.paragraphs)}
t_lookup = {t._tbl: i for i, t in enumerate(doc.tables)}
for n, child in enumerate(doc._element.body.iterchildren()):
    if child.tag == qn("w:p"):
        i = p_lookup[child]
        print(f"{n:03d} P{i:03d}: {doc.paragraphs[i].text}")
    elif child.tag == qn("w:tbl"):
        i = t_lookup[child]
        print(f"{n:03d} T{i:03d}: {doc.tables[i].cell(0,0).text.replace(chr(10), ' / ')}")
    else:
        print(f"{n:03d} {child.tag}")

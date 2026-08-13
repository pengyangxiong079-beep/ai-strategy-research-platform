from __future__ import annotations

import os
from pathlib import Path

import pythoncom
import win32com.client
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\leo\Desktop\ai-strategy-agent")
PPTX_PATH = ROOT / "无生展示课_一般过去时_MemoryPostcard_Revision2.pptx"
POSTCARD_PATH = ROOT / "Memory_Postcard_暑假经历示例.docx"
SCRIPT_PATH = ROOT / "无生展示课_Revision2_完整逐字稿.docx"
PPT_RENDER_DIR = ROOT / "_tmp_revision2_ppt_pages"


def rgb(r, g, b):
    return r + 256 * g + 65536 * b


NAVY = rgb(18, 48, 74)
BLUE = rgb(44, 132, 171)
PALE_BLUE = rgb(224, 243, 248)
CREAM = rgb(250, 247, 239)
SAND = rgb(232, 203, 155)
CORAL = rgb(224, 106, 87)
GREEN = rgb(80, 135, 113)
WHITE = rgb(255, 255, 255)
GRAY = rgb(91, 104, 113)
LIGHT_GRAY = rgb(225, 230, 232)


def set_slide_bg(slide, color=CREAM):
    slide.FollowMasterBackground = False
    slide.Background.Fill.ForeColor.RGB = color
    slide.Background.Fill.Solid()


def add_text(slide, text, left, top, width, height, size=24, color=NAVY,
             bold=False, font="Microsoft YaHei", align=1, valign=1,
             italic=False, margin=4):
    shape = slide.Shapes.AddTextbox(1, left, top, width, height)
    shape.Line.Visible = 0
    shape.Fill.Visible = 0
    tf = shape.TextFrame
    tf.MarginLeft = margin
    tf.MarginRight = margin
    tf.MarginTop = margin
    tf.MarginBottom = margin
    tf.VerticalAnchor = valign
    tr = tf.TextRange
    tr.Text = text
    tr.Font.Name = font
    tr.Font.NameFarEast = "Microsoft YaHei"
    tr.Font.Size = size
    tr.Font.Color.RGB = color
    tr.Font.Bold = -1 if bold else 0
    tr.Font.Italic = -1 if italic else 0
    tr.ParagraphFormat.Alignment = align
    return shape


def add_box(slide, left, top, width, height, fill=WHITE, line=LIGHT_GRAY,
            radius=True, line_width=1.2):
    shape = slide.Shapes.AddShape(5 if radius else 1, left, top, width, height)
    shape.Fill.ForeColor.RGB = fill
    shape.Fill.Solid()
    shape.Line.ForeColor.RGB = line
    shape.Line.Weight = line_width
    return shape


def add_title(slide, title, kicker, number):
    add_text(slide, kicker.upper(), 54, 28, 430, 24, 12, BLUE, True)
    add_text(slide, title, 54, 52, 850, 54, 33, NAVY, True)
    add_text(slide, f"{number:02d}", 902, 34, 30, 24, 11, GRAY, True, align=3)


def add_footer(slide, text="SUMMER MEMORY · PAST TENSE"):
    add_text(slide, text, 54, 510, 300, 20, 10, GRAY, True)
    line = slide.Shapes.AddLine(54, 505, 906, 505)
    line.Line.ForeColor.RGB = LIGHT_GRAY
    line.Line.Weight = 1


def add_notes(slide, text):
    # PowerPoint creates a notes body placeholder automatically.
    try:
        for shape in slide.NotesPage.Shapes:
            if shape.PlaceholderFormat.Type == 2:
                shape.TextFrame.TextRange.Text = text
                return
    except Exception:
        pass


def build_pptx():
    PPT_RENDER_DIR.mkdir(exist_ok=True)
    for old in PPT_RENDER_DIR.glob("*.png"):
        old.unlink()
    pythoncom.CoInitialize()
    app = win32com.client.DispatchEx("PowerPoint.Application")
    app.Visible = True
    app.DisplayAlerts = 0
    pres = app.Presentations.Add()
    pres.PageSetup.SlideWidth = 960
    pres.PageSetup.SlideHeight = 540
    slides = pres.Slides

    # Slide 1
    s = slides.Add(1, 12)
    set_slide_bg(s)
    add_text(s, "A TRIP BACK IN TIME", 54, 50, 500, 28, 14, BLUE, True)
    add_text(s, "Make a Summer\nMemory Postcard", 54, 102, 570, 150, 48, NAVY, True)
    add_text(s, "一般过去时复习 · 九年级", 58, 272, 360, 32, 20, GRAY)
    add_text(s, "Notice → Repair → Ask → Create → Check", 58, 326, 520, 34, 19, CORAL, True)
    card = add_box(s, 650, 82, 245, 350, WHITE, SAND, True, 2)
    card.Rotation = 3
    sun = s.Shapes.AddShape(9, 792, 116, 52, 52)
    sun.Fill.ForeColor.RGB = SAND; sun.Fill.Solid(); sun.Line.Visible = 0
    add_text(s, "QINGDAO", 680, 188, 180, 36, 25, NAVY, True, align=2)
    add_text(s, "last summer", 680, 224, 180, 28, 16, BLUE, False, align=2)
    for y, w in [(286, 165), (306, 190), (326, 145)]:
        line = s.Shapes.AddLine(680, y, 680+w, y)
        line.Line.ForeColor.RGB = PALE_BLUE; line.Line.Weight = 4
    add_text(s, "One memory.\nFour clear sentences.", 686, 354, 180, 54, 17, GREEN, True, align=2)
    add_notes(s, "No external sources. Classroom material authored for this lesson.")

    # Slide 2
    s = slides.Add(2, 12); set_slide_bg(s); add_title(s, "One summer story, one clear goal", "Open the memory", 2)
    add_text(s, "LAST SUMMER", 62, 132, 230, 32, 18, CORAL, True)
    add_text(s, "I went to Qingdao with my parents.", 62, 174, 510, 48, 28, NAVY, True)
    add_text(s, "We walked along the beach, took photos\nand watched the sunset.", 62, 238, 520, 86, 24, GRAY)
    add_text(s, "It was relaxing.", 62, 336, 360, 44, 27, GREEN, True)
    add_text(s, "→", 598, 188, 100, 58, 38, SAND, True, align=2)
    target = add_box(s, 714, 138, 190, 220, WHITE, BLUE, True, 2)
    add_text(s, "TODAY'S TASK", 732, 166, 152, 24, 12, BLUE, True, align=2)
    add_text(s, "Make a\nMemory\nPostcard", 730, 206, 160, 100, 28, NAVY, True, align=2)
    add_text(s, "4 sentences", 748, 315, 124, 26, 16, CORAL, True, align=2)
    add_footer(s)

    # Slide 3
    s = slides.Add(3, 12); set_slide_bg(s); add_title(s, "First, spot the past", "Notice", 3)
    add_text(s, "TIME CLUE", 66, 126, 180, 24, 14, BLUE, True)
    add_text(s, "last summer", 66, 160, 310, 54, 38, CORAL, True)
    line = s.Shapes.AddLine(80, 275, 868, 275); line.Line.ForeColor.RGB = BLUE; line.Line.Weight = 4
    for x in [168, 388, 608, 828]:
        dot = s.Shapes.AddShape(9, x-9, 266, 18, 18); dot.Fill.ForeColor.RGB = WHITE; dot.Fill.Solid(); dot.Line.ForeColor.RGB = BLUE; dot.Line.Weight = 3
    labels = [(120,"go → went"),(330,"walk → walked"),(545,"take → took"),(758,"is → was")]
    for x,txt in labels:
        add_text(s, txt, x, 306, 150, 40, 21, NAVY, True, align=2)
    add_text(s, "Past time clue + past verb", 278, 394, 405, 46, 27, GREEN, True, align=2)
    add_footer(s)

    # Slide 4
    s = slides.Add(4, 12); set_slide_bg(s); add_title(s, "Repair three broken sentences", "Repair", 4)
    rows = [
        ("I go to Qingdao last summer.", "I went to Qingdao last summer.", "past form"),
        ("We didn't took photos.", "We didn't take photos.", "didn't + V"),
        ("Did you watched the sunset?", "Did you watch the sunset?", "Did + S + V"),
    ]
    y = 126
    for i,(wrong,right,rule) in enumerate(rows,1):
        add_text(s, str(i), 62, y+17, 34, 34, 22, CORAL, True, align=2)
        add_text(s, wrong, 112, y, 330, 62, 22, GRAY)
        add_text(s, "→", 452, y+12, 48, 36, 25, BLUE, True, align=2)
        add_text(s, right, 510, y, 330, 62, 22, NAVY, True)
        add_text(s, rule, 770, y+42, 105, 20, 12, GREEN, True, align=3)
        if i < 3:
            sep=s.Shapes.AddLine(62,y+70,890,y+70); sep.Line.ForeColor.RGB=LIGHT_GRAY
        y += 104
    add_text(s, "Find it. Fix it. Say why.", 280, 444, 400, 42, 25, CORAL, True, align=2)
    add_footer(s)

    # Slide 5
    s = slides.Add(5, 12); set_slide_bg(s); add_title(s, "Ask only four questions", "Add details", 5)
    questions = [
        ("WHERE", "Where did you go?", "Qingdao"),
        ("WHO", "Who did you go with?", "my parents"),
        ("WHAT", "What did you do?", "walked · took photos · watched the sunset"),
        ("HOW", "How was it?", "relaxing"),
    ]
    y=126
    for tag,q,a in questions:
        add_text(s, tag, 62, y+5, 92, 32, 15, CORAL, True)
        add_text(s, q, 164, y, 350, 40, 25, NAVY, True)
        add_text(s, a, 530, y, 352, 40, 22, BLUE, False)
        y += 82
    add_text(s, "Listen to the answer. Ask one more detail.", 160, 454, 640, 34, 22, GREEN, True, align=2)
    add_footer(s)

    # Slide 6
    s = slides.Add(6, 12); set_slide_bg(s); add_title(s, "Turn the answers into four sentences", "Create", 6)
    add_text(s, "1", 62, 136, 34, 30, 20, CORAL, True, align=2); add_text(s, "Last summer, I went to Qingdao with my parents.", 112, 128, 770, 46, 23, NAVY, True)
    add_text(s, "2", 62, 204, 34, 30, 20, CORAL, True, align=2); add_text(s, "First, we walked along the beach.", 112, 196, 770, 46, 23, NAVY)
    add_text(s, "3", 62, 272, 34, 30, 20, CORAL, True, align=2); add_text(s, "Then, we took photos and watched the sunset.", 112, 264, 770, 46, 23, NAVY)
    add_text(s, "4", 62, 340, 34, 30, 20, CORAL, True, align=2); add_text(s, "It was relaxing, and I enjoyed time with my family.", 112, 332, 770, 54, 23, NAVY)
    add_box(s, 112, 418, 655, 54, PALE_BLUE, PALE_BLUE, True, 0)
    add_text(s, "Time → Action 1 → Action 2 → Feeling", 126, 428, 625, 34, 22, BLUE, True, align=2)
    add_footer(s)

    # Slide 7
    s = slides.Add(7, 12); set_slide_bg(s); add_title(s, "Check it with PAST", "Improve", 7)
    checks = [
        ("P", "Past time", "Did I say when?", CORAL),
        ("A", "Accurate verbs", "Are my verb forms correct?", BLUE),
        ("S", "Sequence", "Did I use first / then?", GREEN),
        ("T", "True feeling", "Did I share a real feeling?", NAVY),
    ]
    y=124
    for letter,label,q,color in checks:
        circ=s.Shapes.AddShape(9,64,y,54,54); circ.Fill.ForeColor.RGB=color; circ.Fill.Solid(); circ.Line.Visible=0
        add_text(s, letter, 64, y+4, 54, 40, 24, WHITE, True, align=2, valign=3, margin=0)
        add_text(s, label, 142, y+2, 260, 34, 24, NAVY, True)
        add_text(s, q, 424, y+2, 430, 34, 21, GRAY)
        y += 82
    add_text(s, "Give one strength + one helpful suggestion.", 184, 452, 600, 34, 22, CORAL, True, align=2)
    add_footer(s)

    # Slide 8
    s = slides.Add(8, 12); set_slide_bg(s, NAVY)
    add_text(s, "CLOSE THE MEMORY BOX", 56, 40, 430, 24, 13, SAND, True)
    add_text(s, "Every memory can\nhelp us grow.", 56, 82, 520, 118, 43, WHITE, True)
    add_text(s, "Yesterday, Lucy ___ to the library.\nShe didn't ___ TV.  ___ she borrow a book?", 58, 236, 530, 92, 23, WHITE)
    add_text(s, "went · watch · Did", 58, 340, 350, 34, 20, SAND, True)
    add_text(s, "Learning wisely means noticing, checking\nand improving — one small step at a time.", 58, 408, 685, 68, 25, PALE_BLUE, True)
    stamp=s.Shapes.AddShape(9,744,168,132,132); stamp.Fill.ForeColor.RGB=CORAL; stamp.Fill.Solid(); stamp.Line.ForeColor.RGB=SAND; stamp.Line.Weight=3
    add_text(s, "PAST", 748, 196, 124, 40, 29, WHITE, True, align=2)
    add_text(s, "✓", 748, 238, 124, 42, 30, WHITE, True, align=2)
    add_text(s, "Learn wisely.\nRemember deeply.", 706, 330, 208, 65, 18, SAND, True, align=2)
    add_notes(s, "No external sources. Classroom material authored for this lesson.")

    pres.SaveAs(str(PPTX_PATH), 24)
    for slide in slides:
        slide.Export(str(PPT_RENDER_DIR / f"slide-{slide.SlideIndex}.png"), "PNG", 1600, 900)
    pres.Close()
    app.Quit()
    pythoncom.CoUninitialize()
    print(PPTX_PATH)


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, color="D9E1E5", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=160, start=200, bottom=160, end=200):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_run(run, size=11, bold=False, color="12304A", font="Microsoft YaHei", italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(container, text="", size=11, bold=False, color="12304A", align=WD_ALIGN_PARAGRAPH.LEFT,
             before=0, after=4, line=1.1, italic=False):
    p = container.add_paragraph() if hasattr(container, "add_paragraph") else container
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    set_run(r, size, bold, color, italic=italic)
    return p


def configure_doc_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for style_name, size, color in (("Heading 1", 16, "12304A"), ("Heading 2", 13, "2C84AB")):
        st=doc.styles[style_name]; st.font.name="Microsoft YaHei"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)


def build_postcard():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7); sec.page_height = Cm(21)
    sec.top_margin = Cm(1.2); sec.bottom_margin = Cm(1.2); sec.left_margin = Cm(1.4); sec.right_margin = Cm(1.4)
    configure_doc_styles(doc)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3)
    r=p.add_run("SUMMER MEMORY POSTCARD"); set_run(r,24,True,"12304A")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
    r=p.add_run("Last Summer in Qingdao"); set_run(r,13,False,"2C84AB",italic=True)

    table=doc.add_table(rows=1, cols=2); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    table.columns[0].width=Cm(10.7); table.columns[1].width=Cm(14.7)
    left,right=table.rows[0].cells
    left.width=Cm(10.7); right.width=Cm(14.7)
    for c in (left,right): c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c,300,350,300,350); set_cell_border(c,"2C84AB","16")
    set_cell_shading(left,"E0F3F8"); set_cell_shading(right,"FFFFFF")
    left.text=""; right.text=""
    p=left.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("QINGDAO"); set_run(r,23,True,"12304A")
    add_para(left,"SEA · SUNSET · FAMILY",11,True,"E06A57",WD_ALIGN_PARAGRAPH.CENTER,after=12)
    add_para(left,"☀",42,False,"E8CB9B",WD_ALIGN_PARAGRAPH.CENTER,after=8)
    add_para(left,"A small summer memory\nworth keeping",15,True,"508771",WD_ALIGN_PARAGRAPH.CENTER,after=8)
    add_para(left,"TO: My future self",11,False,"5B6871",WD_ALIGN_PARAGRAPH.CENTER)

    p=right.paragraphs[0]; p.paragraph_format.space_after=Pt(8)
    r=p.add_run("Dear friend,"); set_run(r,13,True,"2C84AB")
    sentences=[
        ("Last summer", "Last summer, I went to Qingdao with my parents."),
        ("First", "First, we walked along the beach."),
        ("Then", "Then, we took photos and watched the sunset."),
        ("Feeling", "It was relaxing, and I enjoyed time with my family."),
    ]
    for label,text in sentences:
        p=right.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.1
        r=p.add_run(label+"  "); set_run(r,11,True,"E06A57")
        r=p.add_run(text); set_run(r,13,False,"12304A")
    p=right.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_before=Pt(8)
    r=p.add_run("From, Leo"); set_run(r,12,True,"508771",italic=True)

    doc.add_paragraph()
    check=doc.add_table(rows=1,cols=4); check.alignment=WD_TABLE_ALIGNMENT.CENTER; check.autofit=False
    items=[("P","Past time","Last summer"),("A","Accurate verbs","went · walked · took · watched · was"),("S","Sequence","First · Then"),("T","True feeling","relaxing · family time")]
    widths=[Cm(5.6),Cm(7.2),Cm(5.3),Cm(7.3)]
    for c,w,(letter,label,evidence) in zip(check.rows[0].cells,widths,items):
        c.width=w; set_cell_margins(c,150,160,150,160); set_cell_border(c,"D9E1E5","8")
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(f"{letter} · {label}\n"); set_run(r,10.5,True,"2C84AB")
        r=p.add_run(evidence); set_run(r,9.5,False,"5B6871")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(8)
    r=p.add_run("Learning wisely means noticing, checking and improving — one small step at a time."); set_run(r,11,True,"12304A")
    doc.core_properties.title="Memory Postcard 暑假经历示例"
    doc.save(POSTCARD_PATH)
    print(POSTCARD_PATH)


SCRIPT_SECTIONS = [
    ("Slide 1 · Make a Summer Memory Postcard（0:00—1:00）", [
        ("动作", "微笑，环视；指向标题。"),
        ("教师", "Good morning, boys and girls. Today we are going to take a short trip back in time. We will remember one summer experience and make a Memory Postcard."),
        ("教师", "By the end of the lesson, you only need four clear sentences. We will notice, repair, ask, create and check. Ready?"),
        ("预设学生", "Yes!"),
        ("教师评价", "Great. Let's open the memory box."),
    ]),
    ("Slide 2 · One summer story, one clear goal（1:00—2:10）", [
        ("动作", "指向青岛暑假经历；每问后停顿约2秒。"),
        ("教师", "Look at this summer memory. Last summer, I went to Qingdao with my parents. We walked along the beach, took photos and watched the sunset. It was relaxing."),
        ("教师", "Is this happening now?"),
        ("预设学生", "No. It happened in the past."),
        ("教师", "Exactly. What words tell you the time?"),
        ("预设学生", "Last summer."),
        ("教师评价", "Good noticing. The time clue helps us choose the tense."),
        ("教师", "Our task is simple: use this kind of information to make a four-sentence Summer Memory Postcard."),
    ]),
    ("Slide 3 · First, spot the past（2:10—4:00）", [
        ("动作", "先指 last summer，再沿时间轴依次指四个动词。"),
        ("教师", "First, spot the past. Here is our time clue: last summer. Now look at the verbs."),
        ("教师", "Go changes to...?"),
        ("预设学生", "Went."),
        ("教师", "Walk changes to walked. Take changes to took. Is changes to was."),
        ("教师", "Which two are irregular verbs?"),
        ("预设学生", "Went and took."),
        ("教师评价", "Right. You compared the forms, not only the meanings."),
        ("教师", "So remember this small rule: a past-time clue goes with a past verb. We need the correct past form in an affirmative sentence."),
    ]),
    ("Slide 4 · Repair three broken sentences（4:00—6:30）", [
        ("动作", "教师退到屏幕侧，给学生15秒观察；再从左到右检查。"),
        ("教师", "Three sentences are broken. Be language doctors. Find the problem, fix it and say why. You have fifteen seconds. Begin."),
        ("动作", "静默巡视；15秒后击掌收回。"),
        ("教师", "Number one: I go to Qingdao last summer. What should we change?"),
        ("预设学生", "Go to went."),
        ("教师评价", "Correct. Last summer needs the past form."),
        ("教师", "Number two: We didn't took photos."),
        ("预设学生", "We didn't take photos."),
        ("教师追问", "Why take, not took?"),
        ("预设学生", "Because after didn't, we use the base form."),
        ("教师评价", "Excellent. You repaired it and explained the rule."),
        ("教师", "Number three: Did you watched the sunset?"),
        ("预设学生", "Did you watch the sunset?"),
        ("教师", "Exactly. After Did, use the base form. Find it, fix it, and say why — that is smart grammar learning."),
    ]),
    ("Slide 5 · Ask only four questions（6:30—9:20）", [
        ("动作", "指向四个关键词 Where / Who / What / How。"),
        ("教师", "The grammar is correct, but the memory needs details. We only need four questions."),
        ("教师", "Where did you go?"),
        ("预设学生", "I went to Qingdao."),
        ("教师", "Who did you go with?"),
        ("预设学生", "I went with my parents."),
        ("教师", "What did you do?"),
        ("预设学生", "We walked along the beach, took photos and watched the sunset."),
        ("教师", "How was it?"),
        ("预设学生", "It was relaxing."),
        ("教师评价", "Clear answers. I especially like the detail watched the sunset."),
        ("教师", "Now work in pairs. Student A asks the four questions. Student B answers in full sentences. Then change roles. You have forty seconds. Go."),
        ("动作", "静默走两步，做倾听状；20秒提示 Change roles；40秒收回。"),
        ("教师反馈", "Time's up. I heard a clear question: What did you do? I also heard a useful detail: took photos. Listen to the answer and ask for one more detail."),
    ]),
    ("Slide 6 · Turn the answers into four sentences（9:20—12:20）", [
        ("动作", "四根手指对应四句话。"),
        ("教师", "Now turn the four answers into a postcard. Keep it short and clear."),
        ("教师", "Sentence one gives the time, place and people: Last summer, I went to Qingdao with my parents."),
        ("教师", "Sentence two gives the first action: First, we walked along the beach."),
        ("教师", "Sentence three gives the next action: Then, we took photos and watched the sunset."),
        ("教师", "Sentence four gives the feeling: It was relaxing, and I enjoyed time with my family."),
        ("教师", "The order is easy: Time, Action One, Action Two, Feeling."),
        ("教师指令", "Choose your own summer place and details. Write or say four sentences. You have one minute. The questions on the last slide can help you. Begin."),
        ("动作", "静默巡视约40秒；随后提醒 Check your verbs；60秒收回。"),
        ("教师反馈", "I can see clear time clues. One common point: after First and Then, keep all the main verbs in the past form."),
    ]),
    ("Slide 7 · Check it with PAST（12:20—14:00）", [
        ("动作", "依次用四根手指呈现 P-A-S-T。"),
        ("教师", "Good writers check and improve. Use PAST."),
        ("教师", "P — Past time. Did I say when? A — Accurate verbs. Are my verb forms correct? S — Sequence. Did I use First or Then? T — True feeling. Did I share a real feeling?"),
        ("教师", "Exchange your postcards. Give one strength and one helpful suggestion. Do not only say Good."),
        ("教师示范", "You used a clear time clue. Your verbs are accurate. You can add Then before sentence three. Tell us more about your feeling."),
        ("动作", "停顿20秒，模拟同伴互评。"),
        ("教师反馈", "Thank you. Specific feedback helps a classmate improve. That is also learning wisely."),
    ]),
    ("Slide 8 · Every memory can help us grow（14:00—15:00）", [
        ("动作", "指向出门条，三空各停顿约2秒。"),
        ("教师", "Before we close the memory box, complete the exit ticket. Yesterday, Lucy blank to the library. She didn't blank TV. Blank she borrow a book?"),
        ("预设学生", "Went; watch; Did."),
        ("教师评价", "Excellent. You used the affirmative sentence, the negative sentence and the question correctly."),
        ("教师总结", "Today we used the past tense to make one summer memory clear. We noticed the time, checked the verbs, put events in order and shared a true feeling."),
        ("价值提升", "Learning wisely is not about doing more and more. It means noticing, checking and improving — one small step at a time. Every memory can teach us something, and every reflection can help us grow."),
        ("教师", "Keep your postcard. Keep your memory. Learn wisely, remember deeply, and grow from every experience. Thank you, everyone."),
    ]),
]


def build_script():
    doc=Document(); configure_doc_styles(doc)
    sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.8); sec.left_margin=Cm(2); sec.right_margin=Cm(2)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4)
    r=p.add_run("无生展示课 Revision 2 · 完整逐字稿"); set_run(r,20,True,"12304A")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(12)
    r=p.add_run("一般过去时复习｜核心任务：Summer Memory Postcard｜15分钟"); set_run(r,11,False,"2C84AB")
    table=doc.add_table(rows=2,cols=4); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    vals=[["主题事件","Last Summer in Qingdao","最终产出","4句 Memory Postcard"],["主线步骤","Notice → Repair → Ask → Create → Check","评价工具","PAST 自检"]]
    widths=[Cm(2.5),Cm(6.3),Cm(2.5),Cm(5.4)]
    for row,items in zip(table.rows,vals):
        for cell,w,text in zip(row.cells,widths,items):
            cell.width=w; set_cell_margins(cell,100,120,100,120); set_cell_border(cell,"D9E1E5","6")
            cell.text=""; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(text); set_run(r,9.5, text in {"主题事件","最终产出","主线步骤","评价工具"}, "12304A")
    add_para(doc,"使用说明：教师英语可直接上台使用；“动作、预设学生、教师评价”用于无生课堂模拟。每次提问后保留真实停顿，评价必须引用学生的具体语言。",9.5,False,"5B6871",before=8,after=10,italic=True)
    for heading,blocks in SCRIPT_SECTIONS:
        p=doc.add_paragraph(style="Heading 2"); p.paragraph_format.keep_with_next=True; p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
        r=p.add_run(heading); set_run(r,13,True,"2C84AB")
        for label,text in blocks:
            p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.08
            r=p.add_run(f"【{label}】"); set_run(r,10,True,"E06A57" if label in {"动作","教师指令","教师追问","价值提升"} else "508771")
            r=p.add_run(text); set_run(r,10.5,False,"12304A")
    p=doc.add_paragraph(style="Heading 2"); p.paragraph_format.space_before=Pt(10)
    r=p.add_run("上场前压缩提醒"); set_run(r,13,True,"2C84AB")
    for text in [
        "只保留一个事件：暑假青岛经历；不再加入英语学习方法或第二条主题线。",
        "每个活动只说一次指令：做什么、做到什么程度、用多长时间。",
        "若超时，只减少同伴展示数量，不删“为什么”、PAST自检和出门条。",
        "结尾从语言学习自然过渡到反思成长：Learn wisely, remember deeply, and grow from every experience.",
    ]:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2); r=p.add_run(text); set_run(r,10,False,"12304A")
    doc.core_properties.title="无生展示课 Revision 2 完整逐字稿"
    doc.save(SCRIPT_PATH)
    print(SCRIPT_PATH)


if __name__ == "__main__":
    build_pptx()
    build_postcard()
    build_script()

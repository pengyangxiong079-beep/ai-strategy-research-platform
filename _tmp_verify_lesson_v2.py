from pathlib import Path
from zipfile import ZipFile

from docx import Document


source_dir = Path(r"C:\Users\leo\Desktop\1")
source = min((p for p in source_dir.glob("*.docx") if not p.name.startswith("~$")), key=lambda p: p.stat().st_size)
final = Path(r"C:\Users\leo\Desktop\ai-strategy-agent\无生展示课_v2_一般过去时专业扩充版.docx")
src_doc = Document(source)
out_doc = Document(final)

assert len(src_doc.sections) == len(out_doc.sections) == 1
assert len(src_doc.paragraphs) == len(out_doc.paragraphs) == 49
assert len(src_doc.tables) == len(out_doc.tables) == 8
ss, os = src_doc.sections[0], out_doc.sections[0]
for attr in (
    "page_width", "page_height", "top_margin", "bottom_margin", "left_margin",
    "right_margin", "header_distance", "footer_distance",
):
    assert getattr(ss, attr) == getattr(os, attr), attr
for i, (sp, op) in enumerate(zip(src_doc.paragraphs, out_doc.paragraphs)):
    assert sp.style.name == op.style.name, (i, sp.style.name, op.style.name)
for i, (st, ot) in enumerate(zip(src_doc.tables, out_doc.tables)):
    assert st.style.name == ot.style.name, (i, st.style.name, ot.style.name)
text = "\n".join(p.text for p in out_doc.paragraphs) + "\n" + "\n".join(t.cell(0, 0).text for t in out_doc.tables)
for required in ("didn't + V", "Did + S + V", "weather was", "PAST Self-check", "Exit Ticket", "Memory Postcard"):
    assert required in text, required
with ZipFile(final) as zf:
    bad = zf.testzip()
    assert bad is None, bad
print("PASS")
print("source", source)
print("final", final)
print("paragraphs", len(out_doc.paragraphs), "tables", len(out_doc.tables), "sections", len(out_doc.sections))

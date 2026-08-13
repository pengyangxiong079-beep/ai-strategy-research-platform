from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FONT = "宋体"


def font(run, size=11, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)


def heading(doc, text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    font(r, size, True)
    return p


def para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(22) if indent else None
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    font(r)
    return p


def bullet(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label} ")
    font(r1, 11, True)
    r2 = p.add_run(text)
    font(r2)


def build(output: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(3.0)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("《课堂管理中的“教师”》专业审稿意见")
    font(r, 16, True)

    heading(doc, "一、总体判断")
    para(doc, "原稿具有鲜明的问题意识和较强的一线课堂经验，能够把行动计划、行动向量、活动过渡与干扰处理联系起来，并通过英语课堂对照片段增强可读性。其潜在价值更接近“理论导向的教学反思/教师发展案例”，而非严格意义上的实证研究论文。以原稿状态投稿，建议结论为“大修后再审”；在纠正文献归属、重写方法定位、收缩因果表述并补全参考文献后，可达到面向教师教育、基础教育实践或校本教研类刊物的基本要求。")

    heading(doc, "二、文章结构分解")
    bullet(doc, "1.", "导言提出课堂管理影响有效学习时间，并把教师设定为关键行动者。")
    bullet(doc, "2.", "第一部分以专家—新手研究引出教师专业差异，但原稿误将研究归于杰瑞·布罗菲。")
    bullet(doc, "3.", "第二部分以行动计划和行动向量解释课堂秩序，并用两则课堂处理路径进行对照。")
    bullet(doc, "4.", "第三部分归纳监控、行为期望、程序、过渡、流程和干扰处理六个方面。")
    bullet(doc, "5.", "第四部分转述 Kounin 的监控、重叠、流畅、动力和群体关注五项特征。")
    bullet(doc, "6.", "结论强调优秀课堂管理者需要长期学习和经验积累。")

    heading(doc, "三、主要优点")
    bullet(doc, "1.", "问题来自真实课堂，论述具有实践张力，尤其是干扰如何夺取全班注意这一现象呈现得较为生动。")
    bullet(doc, "2.", "“行动计划—行动向量”的概念能够把备课与现场实施连接起来，是全文最有学术潜力的主线。")
    bullet(doc, "3.", "文章不是单纯罗列纪律技巧，而是注意到监控、过渡、节奏、语言和非语言信号之间的协同。")
    bullet(doc, "4.", "对新手教师计划具体但调整困难、经验教师计划简明但脚本丰富的观察，与专家—新手研究具有较高契合度。")

    heading(doc, "四、必须修改的学术问题")
    bullet(doc, "1. 文献误归属。", "“11位专家、6位新手、每人四节录像课”的研究并非杰瑞·布罗菲2005年的实验，而是 Thiel、Richter 与 Ophardt 2012 年发表的课堂过渡专家—新手研究。该问题属于可影响稿件可信度的硬伤。")
    bullet(doc, "2. 研究设计不清。", "原稿称“进行了一次教学调研”，但没有说明学校范围、教师数量、课次、资料类型、观察维度、分析步骤和伦理处理。若无法补充真实资料，应明确改写为“理论阐释与匿名化典型案例反思”，不能使用“研究证明”“决定”等强因果语言。")
    bullet(doc, "3. 样本标签过度简化。", "把资深且受认可的教师统一归为A组、入职两年内教师归为B组，容易把教龄、声誉、岗位与管理能力混为一谈。建议使用“经验教师/新手教师的典型实践取向”，并声明这不是能力定级。")
    bullet(doc, "4. 论证存在循环。", "原稿一方面用“优秀”作为A组入选条件，另一方面再得出A组管理更优秀，形成选择标准与结论重合。优化稿应把重点放在可观察的行动机制，而非组别高下。")
    bullet(doc, "5. 案例伦理与可识别性。", "原稿直接写出学校和班级编号，且逐字呈现师生冲突，应匿名化并说明对话经过压缩整理。若真实投稿涉及研究数据，还需依目标刊物要求说明知情同意与伦理审查。")
    bullet(doc, "6. 概念关系未澄清。", "“六个方面”与“五个特征”在原稿中存在重复但未解释。建议明确：六个方面是管理任务，五个特征是实施这些任务时表现出的综合能力。")
    bullet(doc, "7. 表述过度。", "“教师个体差异决定课堂效果”“班主任效果显而易见”“完美处理”等说法超出材料支持。宜改为“影响”“在该案例中表现为”“可能与……有关”。")
    bullet(doc, "8. 参考文献缺失。", "正文引用 Doyle、Kounin、Westerman、Emmer 与 Gerwels、Arlin 等作者，却没有参考文献表，也存在 Kounins、Gerwels 拼写和引文页码混乱。投稿前必须统一作者名、年份、题名、出处和格式。")

    heading(doc, "五、语言与格式问题")
    bullet(doc, "1.", "原标题《课堂管理者 “教师”》搭配生硬，建议改为能呈现研究主线的标题。")
    bullet(doc, "2.", "手工目录与正文标题编号不一致，如“2.3”在正文变成“3、”，第三、四部分的分项编号也不统一。")
    bullet(doc, "3.", "正文频繁使用“我们发现”“研究证明”，但证据来源不明确；应区分文献结论、作者观察与分析性推论。")
    bullet(doc, "4.", "存在“有效的管理”“技巧的上”“教学教学环境”等语病和重复，以及中英文标点、空格、大小写不统一。")
    bullet(doc, "5.", "原稿全部使用 Normal 样式，主要依靠手工加粗和缩进，影响后续排版与无障碍导航。若目标刊物提供模板，应在定稿阶段套用其标题层级和引文规范。")

    heading(doc, "六、本次优化策略")
    bullet(doc, "1.", "将文章定位为“理论阐释+匿名化典型案例反思”，不虚构样本量、统计结果或研究程序。")
    bullet(doc, "2.", "纠正专家—新手研究的作者、年份、研究对象与结论，并补充8项核心参考文献。")
    bullet(doc, "3.", "保留原稿“从研究到实践、再回到教师成长”的叙述路径，强化“行动计划—行动向量”主线。")
    bullet(doc, "4.", "把六个管理方面与五个基本特征建立层级关系，增加讨论、实践启示和研究局限。")
    bullet(doc, "5.", "匿名化学校、班级和师生信息，压缩对话，改为机制分析，避免将个案变成对教师或学生的价值评判。")
    bullet(doc, "6.", "延续原稿楷体、A4、黑白、正文两字符缩进与简洁风格，只对标题、摘要、编号和段落节奏作必要规范化。")

    heading(doc, "七、发表价值与建议")
    para(doc, "优化后稿件的主要价值在于：以较少进入一线教师写作的“行动向量”概念解释课堂干扰、过渡与流程保护，并把专家—新手差异转化为可训练的观察、解释和决策能力。该稿适合投向教师教育、课堂教学、基础教育实践或校本研修类刊物。若拟投教育学核心或高水平实证期刊，仍需另行补充可复核的研究设计、系统数据、伦理说明和更完整的国内外文献综述；仅凭当前材料不宜包装为实证研究。")

    doc.core_properties.title = "《课堂管理中的“教师”》专业审稿意见"
    doc.core_properties.subject = "学术审稿与修订建议"
    doc.save(output)


if __name__ == "__main__":
    build(Path(sys.argv[1]))

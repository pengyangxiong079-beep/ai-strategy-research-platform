from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from article_content import ABSTRACT, KEYWORDS, REFERENCES, SECTIONS, TITLE


FONT = "楷体"


def set_font(run, size=12, bold=None, italic=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)


def keep_lines(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    for tag in ("keepLines", "widowControl"):
        el = ppr.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            ppr.append(el)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = True
    keep_lines(p)
    r = p.add_run(text)
    set_font(r, 12, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.35
    keep_lines(p)
    r = p.add_run(text)
    set_font(r, 12, bold=False)
    return p


def add_labeled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    keep_lines(p)
    r1 = p.add_run(label)
    set_font(r1, 12, bold=True)
    r2 = p.add_run(text)
    set_font(r2, 12, bold=False)
    return p


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    keep_lines(p)
    r = p.add_run(text)
    set_font(r, 10.5, bold=False)
    return p


def clear_body_preserving_section(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_core(doc):
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "课堂管理与教师专业发展"
    doc.core_properties.keywords = "课堂管理, 教师专业差异, 专家新手教师, 行动向量"
    doc.core_properties.comments = "在保留原稿楷体、A4页面与简洁风格基础上的学术优化稿。"


def build(source: Path, output: Path):
    doc = Document(source)
    clear_body_preserving_section(doc)
    set_core(doc)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    title = doc.add_paragraph()
    title.style = normal
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0
    title.paragraph_format.keep_with_next = True
    tr = title.add_run(TITLE)
    set_font(tr, 16, bold=True)

    add_labeled_paragraph(doc, "摘  要：", ABSTRACT)
    add_labeled_paragraph(doc, "关键词：", KEYWORDS)

    for heading, paragraphs in SECTIONS:
        level = 1 if heading in {"导言", "结论"} or heading[:1] in "一二三四五六七八九十" else 2
        add_heading(doc, heading, level)
        for text in paragraphs:
            add_body(doc, text)

    add_heading(doc, "参考文献", 1)
    for ref in REFERENCES:
        add_reference(doc, ref)

    doc.save(output)


if __name__ == "__main__":
    build(Path(sys.argv[1]), Path(sys.argv[2]))

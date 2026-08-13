from pathlib import Path
from zipfile import ZipFile
import re

from docx import Document


root = Path(r"C:\Users\leo\Desktop\ai-strategy-agent")
pptx = root / "无生展示课_一般过去时_MemoryPostcard_Revision2.pptx"
postcard = root / "Memory_Postcard_暑假经历示例.docx"
script = root / "无生展示课_Revision2_完整逐字稿.docx"

for path in (pptx, postcard, script):
    assert path.exists() and path.stat().st_size > 10000, path
    with ZipFile(path) as archive:
        assert archive.testzip() is None, path

with ZipFile(pptx) as archive:
    slide_parts = [n for n in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", n)]
    assert len(slide_parts) == 8, len(slide_parts)
    deck_text = "".join(archive.read(n).decode("utf-8", errors="ignore") for n in slide_parts)
    for required in ("Memory Postcard", "last summer", "didn't take", "Did you watch", "PAST", "Learning wisely"):
        assert required in deck_text, required

post_doc = Document(postcard)
post_text = "\n".join(p.text for p in post_doc.paragraphs) + "\n" + "\n".join(c.text for t in post_doc.tables for row in t.rows for c in row.cells)
for required in ("Last summer, I went to Qingdao", "First", "Then", "P · Past time", "A · Accurate verbs", "S · Sequence", "T · True feeling", "Learning wisely"):
    assert required in post_text, required

script_doc = Document(script)
script_text = "\n".join(p.text for p in script_doc.paragraphs) + "\n" + "\n".join(c.text for t in script_doc.tables for row in t.rows for c in row.cells)
for required in ("Slide 1", "Slide 8", "didn't take", "Did you watch", "PAST", "Learn wisely, remember deeply"):
    assert required in script_text, required

print("PASS")
print("slides", len(slide_parts))
print("postcard_tables", len(post_doc.tables))
print("script_paragraphs", len(script_doc.paragraphs))
